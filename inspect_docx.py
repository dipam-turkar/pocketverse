
import docx

def inspect_docx(file_path):
    doc = docx.Document(file_path)
    print(f"Total Paragraphs: {len(doc.paragraphs)}")
    print("-" * 20)
    for i, para in enumerate(doc.paragraphs[:50]):
        if para.text.strip():
            print(f"Para {i}: {para.text[:100]}")
    
    print("\n" * 2)
    print(f"Total Tables: {len(doc.tables)}")
    for i, table in enumerate(doc.tables[:5]):
        print(f"Table {i} rows: {len(table.rows)}")
        for row in table.rows[:3]:
            print(" | ".join([cell.text.strip() for cell in row.cells]))
            
if __name__ == "__main__":
    inspect_docx("PromoCanon/Saving_Nora_CHARACTER BIBLE.docx")
